from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Inches

from flask import send_file
from flask_login import login_required

from app.models import Item, RequestedItem, User
from app.utils.is_admin import is_admin
from . import module


# creating a report
@module.route('/report', methods=['GET'])
@login_required
@is_admin
def get_report():
    # Получаем все предметы из базы данных
    items = Item.query.all()

    # Создаем новый документ
    doc = Document()
    doc.add_heading('Отчет об инвентаре', level=1)

    # Добавление даты создания отчета
    report_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_paragraph(f'Дата создания отчета: {report_date}')
    doc.add_paragraph()  # Пустая строка для разделения

    # Добавляем таблицу в документ
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    # Устанавливаем ширину столбцов
    for i in range(len(table.columns)):
        if i == 0:
            table.columns[i].width = Inches(3)
        elif i == 1:
            table.columns[i].width = Inches(2)
        elif i == 2:
            table.columns[i].width = Inches(2)
        elif i == 3:
            table.columns[i].width = Inches(2)
        elif i == 4:
            table.columns[i].width = Inches(6)

    # Добавляем заголовок таблицы
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Название предмета'
    hdr_cells[1].text = 'Общее количество'
    hdr_cells[2].text = 'Доступное количество'
    hdr_cells[3].text = 'Состояние'
    hdr_cells[4].text = 'Используется (пользователь и количество)'

    # Добавляем строки с данными
    for item in items:
        row_cells = table.add_row().cells
        row_cells[0].text = item.name
        row_cells[1].text = str(item.total_quantity)
        row_cells[2].text = str(item.available_quantity)
        row_cells[3].text = item.state.value

        # Получаем пользователей, использующих данный предмет
        users_using_item = RequestedItem.query.filter_by(item_id=item.id).all()
        if users_using_item:
            user_details = []
            for user in users_using_item:
                user_name = User.query.get(user.user_id).name
                user_quantity = user.quantity
                user_details.append(f"{user_name} ({user_quantity})")
            row_cells[4].text = ', '.join(user_details)
        else:
            row_cells[4].text = 'Нет'

    # Сохраняем документ в объект BytesIO
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)

    # Отправляем файл пользователю
    return send_file(byte_io, as_attachment=True, download_name='отчет_об_инвентаре.docx',
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')